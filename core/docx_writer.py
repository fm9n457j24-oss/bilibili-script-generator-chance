# -*- coding: utf-8 -*-
"""
DOCX 文档生成模块
按照用户上传的「混剪脚本-模板」格式生成 Word 文档。
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn


# 表格列定义（与模板一致）
COLUMNS = ["镜号", "后期", "画面参考", "字幕/台词", "备注"]
# 列宽（英寸）—— 总和约 9.3 英寸适配 A4 横向
COL_WIDTHS = [0.6, 1.8, 3.0, 2.5, 1.4]


def generate_docx(script_data: dict, output_path: str) -> str:
    """
    根据脚本数据生成 Word 文档。

    Args:
        script_data: 脚本数据 dict，包含:
            - title: 脚本标题
            - duration: 时长
            - orientation: 横版/竖版
            - style_reference: 剪辑风格参考
            - rows: [{shot_number, post_production, visual_reference,
                      subtitle_dialogue, notes}, ...]
        output_path: 输出文件路径

    Returns:
        输出文件路径
    """
    doc = Document()

    # ---------------------------------------------------------------- #
    #  页面设置：A4 横向
    # ---------------------------------------------------------------- #
    section = doc.sections[0]
    section.orientation = 1  # WD_ORIENT.LANDSCAPE
    # 交换宽高
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ---------------------------------------------------------------- #
    #  默认字体设置（支持中文）
    # ---------------------------------------------------------------- #
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    # 设置东亚字体
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ---------------------------------------------------------------- #
    #  标题
    # ---------------------------------------------------------------- #
    title_text = script_data.get("title", "混剪脚本")
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"《{title_text}》-混剪脚本")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "微软雅黑"
    title_run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ---------------------------------------------------------------- #
    #  基本信息
    # ---------------------------------------------------------------- #
    duration = script_data.get("duration", "2-3分钟")
    orientation = script_data.get("orientation", "横版")
    style_ref = script_data.get("style_reference", "")
    hook_type = script_data.get("hook_type", "")
    script_structure = script_data.get("script_structure", "")

    info_lines = [
        f"时长：{duration}",
        f"横/竖板：{orientation}",
        f"剪辑风格参考：{style_ref}",
    ]
    if hook_type:
        info_lines.append(f"开头钩子类型：{hook_type}")
    if script_structure:
        info_lines.append(f"脚本结构：{script_structure}")
    for line in info_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.name = "微软雅黑"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        p.paragraph_format.space_after = Pt(2)

    # 空行
    doc.add_paragraph()

    # ---------------------------------------------------------------- #
    #  分镜表格
    # ---------------------------------------------------------------- #
    rows_data = script_data.get("rows", [])

    table = doc.add_table(rows=1 + len(rows_data), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 设置列宽
    for i, width in enumerate(COL_WIDTHS):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)

    # 表头
    header_cells = table.rows[0].cells
    for i, col_name in enumerate(COLUMNS):
        cell = header_cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col_name)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "微软雅黑"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        # 表头背景色
        _set_cell_bg(cell, "4472C4")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 数据行
    for row_idx, row_data in enumerate(rows_data):
        row_cells = table.rows[row_idx + 1].cells

        values = [
            str(row_data.get("shot_number", row_idx + 1)),
            row_data.get("post_production", ""),
            row_data.get("visual_reference", ""),
            row_data.get("subtitle_dialogue", ""),
            row_data.get("notes", ""),
        ]

        for i, val in enumerate(values):
            cell = row_cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            if i == 0:  # 镜号居中
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = p.add_run(val)
            run.font.size = Pt(10.5)
            run.font.name = "微软雅黑"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 隔行变色
            if row_idx % 2 == 1:
                _set_cell_bg(cell, "EDF2FA")

    # ---------------------------------------------------------------- #
    #  保存
    # ---------------------------------------------------------------- #
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def _set_cell_bg(cell, color_hex: str):
    """设置单元格背景色"""
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)
